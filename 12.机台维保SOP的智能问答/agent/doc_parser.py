"""
SOP 文档解析器 — 支持 .md/.txt/.pdf/.docx
提取: 元数据、章节结构、步骤列表、标准值、工具清单、安全项
"""
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple


class SOPDocumentParser:
    """SOP 文档解析器"""

    SUPPORTED_EXT = {'.md', '.txt', '.pdf', '.docx'}

    @staticmethod
    def parse(filepath: str) -> Dict:
        """
        解析 SOP 文档, 返回结构化数据

        Returns:
            {
                "filepath": str,
                "metadata": {编号, 版本, 设备, 维护类型, 工时, 人员, ...},
                "sections": [{heading, level, number, content, is_safety, tools, standards}],
                "safety_items": [{section, description, requirement}],
                "tool_list": [str],
                "standard_values": [{param, value, section}],
                "step_count": int,
                "raw_text": str,
            }
        """
        fp = Path(filepath)
        if not fp.exists():
            raise FileNotFoundError(f"SOP 文件不存在: {filepath}")

        ext = fp.suffix.lower()

        if ext == '.md' or ext == '.txt':
            raw = fp.read_text(encoding="utf-8")
        elif ext == '.pdf':
            raw = SOPDocumentParser._parse_pdf(fp)
        elif ext == '.docx':
            raw = SOPDocumentParser._parse_docx(fp)
        else:
            raise ValueError(f"不支持的格式: {ext} (支持: .md/.txt/.pdf/.docx)")

        metadata = SOPDocumentParser._extract_metadata(raw)
        sections = SOPDocumentParser._extract_sections(raw)
        safety_items = SOPDocumentParser._extract_safety(raw)
        tools = SOPDocumentParser._extract_tools(raw)
        standards = SOPDocumentParser._extract_standards(raw, sections)

        return {
            "filepath": str(fp),
            "metadata": metadata,
            "sections": sections,
            "safety_items": safety_items,
            "tool_list": tools,
            "standard_values": standards,
            "step_count": len([s for s in sections if s.get("number")]),
            "raw_text": raw,
        }

    @staticmethod
    def _parse_pdf(filepath: Path) -> str:
        """PDF 解析 (pdfplumber)"""
        try:
            import pdfplumber
            parts = []
            with pdfplumber.open(str(filepath)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        parts.append(text)
                    # Also extract tables
                    for table in page.extract_tables():
                        for row in table:
                            parts.append(' | '.join(str(c) for c in row if c))
            return '\n'.join(parts)
        except ImportError:
            raise ImportError("PDF 解析需要 pdfplumber: pip install pdfplumber")

    @staticmethod
    def _parse_docx(filepath: Path) -> str:
        """DOCX 解析 (python-docx)"""
        try:
            from docx import Document
            doc = Document(str(filepath))
            parts = []
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading ', '')
                    prefix = '#' * max(1, int(level)) if level.isdigit() else '#'
                    parts.append(f"{prefix} {para.text}")
                else:
                    parts.append(para.text)
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    parts.append(' | '.join(cell.text for cell in row.cells))
            return '\n'.join(parts)
        except ImportError:
            raise ImportError("DOCX 解析需要 python-docx: pip install python-docx")

    @staticmethod
    def _extract_metadata(text: str) -> Dict:
        """从 SOP 文档头部提取元数据"""
        meta = {}
        meta_patterns = {
            "doc_id": r'文件编号[：:]\s*(.+)',
            "version": r'版本[：:]\s*(.+)',
            "equipment": r'设备[：:]\s*(.+)',
            "maintenance_type": r'维护类型[：:]\s*(.+)',
            "estimated_time": r'预计工时[：:]\s*(.+)',
            "personnel": r'所需人员[：:]\s*(.+)',
            "effective_date": r'生效日期[：:]\s*(.+)',
        }
        for key, pattern in meta_patterns.items():
            m = re.search(pattern, text)
            if m:
                meta[key] = m.group(1).strip()
        return meta

    @staticmethod
    def _extract_sections(text: str) -> List[Dict]:
        """提取章节结构"""
        sections = []
        heading_re = re.compile(
            r'^(#{1,4})[\s]*(\d+(?:\.\d+)*)?[\s]*(.+)?$', re.MULTILINE
        )
        lines = text.split('\n')

        cur = {"heading": "前言", "level": 0, "number": "", "content": "",
               "is_safety": False, "tools": [], "standards": []}
        buffer_lines = []

        for line in lines:
            m = heading_re.match(line)
            if m:
                if buffer_lines:
                    cur["content"] = '\n'.join(buffer_lines).strip()
                    # Post-process: detect safety, tools, standards
                    SOPDocumentParser._tag_section(cur)
                    sections.append(cur)

                heading_text = (m.group(3) or m.group(2) or line).strip()
                cur = {
                    "heading": heading_text,
                    "level": len(m.group(1)),
                    "number": m.group(2) or "",
                    "content": "",
                    "is_safety": False,
                    "tools": [],
                    "standards": [],
                }
                buffer_lines = []
            else:
                buffer_lines.append(line)

        # flush last
        if buffer_lines or cur.get("heading"):
            cur["content"] = '\n'.join(buffer_lines).strip()
            SOPDocumentParser._tag_section(cur)
            sections.append(cur)

        return sections

    @staticmethod
    def _tag_section(section: Dict):
        """标注章节: 安全/工具/标准值"""
        combined = (section["heading"] + " " + section["content"]).lower()

        # 安全标注
        safety_kw = ["安全", "emo", "ppe", "lockout", "tagout", "loto", "泄漏", "有害气体",
                     "警示", "警告", "危险", "通风", "上锁", "挂牌"]
        if any(kw in combined for kw in safety_kw):
            section["is_safety"] = True

        # 工具识别
        tool_pattern = re.compile(
            r'(扭矩扳手|氦检漏仪|真空吸尘器|Scotch[\s-]*Brite|无尘布|IPA|DI Water|'
            r'N2|塞尺|热像仪|网络分析仪|O[\s-]*ring pick|真空脂|Krytox|'
            r'torque wrench|helium leak detector|feeler gauge|thermal camera|'
            r'network analyzer)',
            re.IGNORECASE
        )
        section["tools"] = list(set(
            m.group(0).strip() for m in tool_pattern.finditer(combined)
        ))

        # 标准值识别
        spec_pattern = re.compile(
            r'(?P<param>(?:温度|温差|压力|漏率|泄漏率|电阻|VSWR|驻波比|RF hours|'
            r'厚度|间隙|Etch Rate|CD|NU%|电流|油位))[^\d]*(?P<value>\d+(?:\.\d+)?\s*'
            r'(?:°C|ppm|Torr|mTorr|MΩ|mm|μm|nm|sccm|mbar|L/s|%|小时|min))?',
            re.IGNORECASE
        )
        for m in spec_pattern.finditer(combined):
            section["standards"].append({
                "param": m.group("param"),
                "value": m.group("value") or "见原文",
            })

    @staticmethod
    def _extract_safety(text: str) -> List[Dict]:
        """提取所有安全相关条目"""
        items = []
        safety_pattern = re.compile(
            r'(?:⚠️|安全|EMO|PPE|Lockout|Tagout|LOTO|上锁|挂牌|警示|警告|危险|'
            r'有害气体|泄漏|通风|禁止操作|emergency)',
            re.IGNORECASE
        )
        for line in text.split('\n'):
            if safety_pattern.search(line):
                items.append({
                    "line": line.strip(),
                    "has_emo": bool(re.search(r'EMO', line, re.IGNORECASE)),
                    "has_ppe": bool(re.search(r'PPE|手套|护目镜|安全帽|防静电', line)),
                    "has_vent": bool(re.search(r'通风', line)),
                })
        return items

    @staticmethod
    def _extract_tools(text: str) -> List[str]:
        """提取工具和物料清单"""
        tool_set = set()
        patterns = [
            r'扭矩扳手', r'氦检漏仪', r'真空吸尘器', r'Scotch[\s-]*Brite',
            r'无尘布', r'IPA', r'DI Water', r'去离子水', r'[Nn]2\b(?!\w)', r'氮气',
            r'塞尺', r'热像仪', r'网络分析仪', r'O[\s-]*ring pick',
            r'真空脂', r'Krytox', r'专用吊具', r'专用清洁剂', r'Clean[\s-]*100',
            r'Monitor Wafer', r'Dummy Wafer', r'扭矩', r'扳手',
        ]
        for pat in patterns:
            for m in re.finditer(pat, text, re.IGNORECASE):
                tool_set.add(m.group(0).strip())
        return sorted(tool_set)

    @staticmethod
    def _extract_standards(text: str, sections: List[Dict]) -> List[Dict]:
        """提取所有标准值/阈值"""
        standards = []
        spec_pattern = re.compile(
            r'(?P<param>温度|温差|压力|漏率|电阻|VSWR|驻波比|厚度|间隙|'
            r'Etch Rate|CD|NU%|电流|RF hours|清洁度)'
            r'[^\d<>]*?(?P<op>[<>≤≥]?\s*)'
            r'(?P<value>\d+(?:\.\d+)?\s*(?:°C|ppm|Torr|mTorr|MΩ|mm|μm|nm|sccm|mbar|L/s|%|小时|min|片))',
            re.IGNORECASE
        )
        for m in spec_pattern.finditer(text):
            standards.append({
                "param": m.group("param").strip(),
                "value": f"{m.group('op')} {m.group('value')}".strip(),
                "raw": m.group(0),
            })
        return standards
